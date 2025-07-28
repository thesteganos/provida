from docx import Document
from docx.shared import Inches
from app.models.research_models import FinalReport
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)

class DOCXExporter:
    """
    Classe para exportar relatórios de pesquisa para DOCX.
    """
    def __init__(self):
        """
        Inicializa o DOCXExporter, criando um novo documento DOCX.
        """
        self.document = Document()

    def add_title(self, title: str):
        """
        Adiciona um título principal ao documento.

        Args:
            title (str): O texto do título.
        """
        self.document.add_heading(title, level=1)

    def add_section_title(self, title: str):
        """
        Adiciona um título de seção ao documento.

        Args:
            title (str): O texto do título da seção.
        """
        self.document.add_heading(title, level=2)

    def add_text(self, text: str):
        """
        Adiciona um parágrafo de texto ao documento.

        Args:
            text (str): O texto a ser adicionado.
        """
        self.document.add_paragraph(text)

    def add_citations(self, citations: List[Dict[str, Any]]):
        """
        Adiciona uma seção de citações ao documento.

        Args:
            citations (List[Dict[str, Any]]): Uma lista de dicionários de citação.
        """
        self.add_section_title("Citações Utilizadas")
        for c in citations:
            citation_text = f"- [ID: {c.id}] {c.sentence_in_summary}"
            self.document.add_paragraph(citation_text)

    def export_report(self, final_report: FinalReport, output_filename: str = "report.docx"):
        # Ensure final_report is a FinalReport instance
        if not isinstance(final_report, FinalReport):
            logger.error("Invalid input: final_report must be an instance of FinalReport.")
            return False
        """
        Exporta o relatório final para um arquivo DOCX.

        Args:
            final_report (FinalReport): O objeto FinalReport contendo os dados do relatório.
            output_filename (str): O nome do arquivo de saída (padrão: "report.docx").

        Returns:
            bool: True se o relatório foi exportado com sucesso, False caso contrário.
        """
        logger.info(f"Exportando relatório para DOCX: {output_filename}")
        
        self.add_title("Relatório de Pesquisa Pró-Vida")
        self.add_section_title(f"Tópico: {final_report.research_question}")
        self.add_text(f"Data de Geração: {final_report.generation_date}")
        self.document.add_paragraph("") # Add a blank line

        self.add_section_title("Resumo Final")
        self.add_text(final_report.summary)

        citations = final_report.citations_used
        if citations:
            self.add_citations(citations)

        try:
            self.document.save(output_filename)
            logger.info(f"Relatório DOCX salvo em {output_filename}")
            return True
        except Exception as e:
            logger.error(f"Erro ao salvar o relatório DOCX: {e}", exc_info=True)
            return False
